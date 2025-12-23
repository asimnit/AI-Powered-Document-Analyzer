"""
Word Document Parser

Extracts text from Microsoft Word documents (.docx) using python-docx.
"""

import logging
from app.services.parsers.base import DocumentParser, ParserResult

logger = logging.getLogger(__name__)


class WordParser(DocumentParser):
    """
    Microsoft Word (.docx) parser
    
    Extracts text from Word documents including:
    - Paragraphs
    - Tables
    - Headers/footers
    """
    
    def parse(self, file_path: str) -> ParserResult:
        """
        Extract text from Word document
        
        Args:
            file_path: Path to .docx file
            
        Returns:
            ParserResult with extracted text
        """
        if not self._validate_file(file_path):
            return ParserResult(
                text="",
                error=f"File not found or not readable: {file_path}"
            )
        
        try:
            from docx import Document
            
            doc = Document(file_path)
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = "\t".join(cell.text for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            full_text = "\n".join(text_parts)
            
            # Estimate page count (rough: 500 words per page)
            word_count = len(full_text.split())
            page_count = max(1, word_count // 500)
            
            return ParserResult(
                text=full_text,
                page_count=page_count,
                word_count=word_count,
                metadata={"parser": "python-docx"}
            )
            
        except Exception as e:
            error_msg = f"Failed to parse Word document: {str(e)}"
            logger.error(error_msg)
            return ParserResult(text="", error=error_msg)
    
    def supports(self, file_extension: str) -> bool:
        """Check if file extension is .docx"""
        return file_extension.lower() in [".docx", ".doc"]
